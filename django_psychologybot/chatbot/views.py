from django.shortcuts import render, redirect
from django.http import JsonResponse, HttpResponse
from openai import OpenAI, AssistantEventHandler
from typing_extensions import override
from . import creds
from django_psychologybot import settings
from docx import Document
<<<<<<< Updated upstream
import os
import io
=======
from django.views.decorators.csrf import csrf_exempt
from .video_analysis import analysis
from threading import Lock
from .models import Report
import os
import io
import subprocess
>>>>>>> Stashed changes


assistant_id = creds.assistant_id

client = OpenAI(
    api_key=creds.api_key
)

<<<<<<< Updated upstream
=======

@csrf_exempt
def upload_video(request):
    if request.method == 'POST' and request.FILES.get('video'):
        video_file = request.FILES['video'];
        video_file_path = os.path.join('static', 'video', video_file.name)

        with open(video_file_path, 'wb') as file:
            for chunk in video_file.chunks():
                file.write(chunk)

        output_video_path = f'{video_file.name}_output.mp4'

        convertion_command = ['ffmpeg', '-i', video_file_path, '-c:v', 
            'libx264', '-c:a', 'aac', output_video_path]
        
        subprocess.run(convertion_command)

        given_emotions = analysis.extract_faces(output_video_path, os.path.join('static', 'output_dir'))

        return JsonResponse({'emotions': given_emotions})


thread_lock = Lock()


>>>>>>> Stashed changes
class EventHandler(AssistantEventHandler):
    def __init__(self):
        super().__init__()
        self.response = ""
        
    @override
    def on_text_delta(self, delta, snapshot):
        self.response += delta.value


    def get_response(self):
        return self.response

# thread = client.beta.threads.create()

QUESTIONNAIRE_LABEL_RU = 'Анкетирование завершено!'
QUESTIONNAIRE_LABEL_KZ = 'Сауалдау аяқталды!'


def generate_word(content):
    """
    Генерирует Word-документ из текста.
    """
    # Создаём объект Word-документа
    document = Document()
    
    # Добавляем заголовок
    document.add_heading("Отчёт ассистента", level=1)
    
    # Добавляем основной текст (контент)
    document.add_paragraph(content)
    
    # Сохраняем документ в поток
    word_file = io.BytesIO()
<<<<<<< Updated upstream
    document.save(word_file)
=======
    document.save(word_file)    
>>>>>>> Stashed changes
    word_file.seek(0)


    return word_file


<<<<<<< Updated upstream
def ask_openai_with_assistant(message, thread_id):
    # Отправляем сообщение пользователю
    client.beta.threads.messages.create(
        thread_id=thread_id,
        role="user",
        content=message
    )
=======
def ask_openai_with_assistant(message, thread_id, report_id):
    # Отправляем сообщение пользователю
    with thread_lock:
        client.beta.threads.messages.create(
            thread_id=thread_id,
            role="user",
            content=message
        )
>>>>>>> Stashed changes

    event_handler = EventHandler()

    with client.beta.threads.runs.stream(
        thread_id=thread_id,
        assistant_id=assistant_id,
        event_handler=event_handler,
    ) as stream:
        stream.until_done()

    # return event_handler.get_response()
    assistant_response = event_handler.get_response()

    if QUESTIONNAIRE_LABEL_RU in assistant_response or QUESTIONNAIRE_LABEL_KZ in assistant_response:
        word_file = generate_word(assistant_response)
<<<<<<< Updated upstream
        word_file_path = os.path.join(settings.STATIC_ROOT, 'report.docx')
        word_file_url = 'http://127.0.0.1:8000/static/report.docx'

        with open(word_file_path, 'wb') as f:
            f.write(word_file.getvalue())
=======
        try:
            # Получаем отчет по ID
            report = Report.objects.get(id=report_id)
            full_name = report.full_name.replace(" ", "_")  # Заменяем пробелы на подчеркивания

            # Создаем уникальное имя файла с ФИО
            word_file_name = f"{full_name}_report.docx"
            word_file_path = os.path.join('static', 'reports', word_file_name)
            word_file_url = f'http://127.0.0.1:8000/static/reports/{word_file_name}'

            os.makedirs(os.path.dirname(word_file_path), exist_ok=True)
            
            # Сохраняем файл на диск
            with open(word_file_path, 'wb') as f:
                f.write(word_file.getvalue())

            # Обновляем запись в модели Report с путем к файлу
            report.file = word_file_path  # Сохраняем путь к файлу в поле file
            report.save()

        except Report.DoesNotExist:
            # Обработка случая, если отчет с таким id не найден
            pass
        # word_file_path = os.path.join(settings.STATIC_ROOT, 'report.docx')
        # word_file_url = 'https://s983114.srvape.com/static/report.docx'

        # with open(word_file_path, 'wb') as f:
        #     f.write(word_file.getvalue())
>>>>>>> Stashed changes

        # Перенаправляем на URL для скачивания
        # return JsonResponse({
        #     'message': 'Анкетирование завершено. Отчёт готов для скачивания.',
        #     'file_url': os.path.join(settings.MEDIA_URL, 'assistant_report.pdf')
        # })
        return assistant_response, word_file_url

    return assistant_response, None
    # return JsonResponse({'message': assistant_response})


def login(request):
    if request.method == 'POST':
        username = request.POST.get('username', '').strip()
        number = request.POST.get('number', '').strip()
        language = request.POST.get('language', '').strip()

        if username and number and language:
<<<<<<< Updated upstream
=======
            report = Report.objects.create(
                full_name=username,
                military_number=number,
                language=language,
                file=None
            )

            request.session['report_id'] = report.id
>>>>>>> Stashed changes
            thread = client.beta.threads.create()
            request.session['thread_id'] = thread.id
            request.session['message'] = f"Меня зовут {username}, мой номер {number} и я хочу разговаривать на {language}."
            return redirect('chatbot')
        else:
            error_message = 'Пожалуйста, заполните все поля.'
            return render(request, 'login.html', {'error_message': error_message})
    else:
        return render(request, 'login.html')


<<<<<<< Updated upstream
def video_callback(request):
    return HttpResponse('Callback')



=======
>>>>>>> Stashed changes
def chatbot(request):
    if request.method == 'POST':
        message = request.POST.get('message')

        thread_id = request.session.get('thread_id')
<<<<<<< Updated upstream
=======
        report_id = request.session.get('report_id')
>>>>>>> Stashed changes

        if not thread_id:
            return redirect('login')

<<<<<<< Updated upstream
        response, file_url = ask_openai_with_assistant(message, thread_id)
=======
        response, file_url = ask_openai_with_assistant(message, thread_id, report_id)
>>>>>>> Stashed changes
        return JsonResponse({'message': message, 'response': response, 'fileURL': file_url})

    message = request.session.get('message', 'Данные отсутствуют.')
    thread_id = request.session.get('thread_id')
<<<<<<< Updated upstream
=======
    report_id = request.session.get('report_id')
>>>>>>> Stashed changes

    if not thread_id:
        return redirect('login')

<<<<<<< Updated upstream
    response, _ = ask_openai_with_assistant(message, thread_id)
=======
    response, _ = ask_openai_with_assistant(message, thread_id, report_id)
>>>>>>> Stashed changes
    return render(request, 'chatbot.html', {'message': message, 'response': response})
